import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import os

class RelatorioObrasApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Gerador de Relatório Fotográfico v1.2")
        self.geometry("550x750") # Tela mais compacta
        self.minsize(500, 700)
        ctk.set_appearance_mode("dark")

        self.lista_fotos = []
        self.caminho_logo_final = None
        self.diretorio_logos = "Logos"
        self.img_preview_referencia = None 

        if not os.path.exists(self.diretorio_logos):
            os.makedirs(self.diretorio_logos)

        # --- Frame Principal ---
        self.main_container = ctk.CTkScrollableFrame(self, fg_color="transparent")
        self.main_container.pack(fill="both", expand=True, padx=5, pady=5)

        ctk.CTkLabel(self.main_container, text="Relatório Fotográfico Profissional", font=("Arial", 18, "bold")).pack(pady=10)

        # Entradas de texto
        self.entry_obra = ctk.CTkEntry(self.main_container, placeholder_text="Nome da Obra / Cliente", width=420)
        self.entry_obra.pack(pady=2)
        self.entry_responsavel = ctk.CTkEntry(self.main_container, placeholder_text="Responsável Técnico (Nome/CREA)", width=420)
        self.entry_responsavel.pack(pady=2)
        self.entry_data = ctk.CTkEntry(self.main_container, placeholder_text="Data da Visita (ex: 12/03/2031)", width=420)
        self.entry_data.pack(pady=2)

        # --- Seção de Logomarca (Compactada) ---
        self.frame_logo = ctk.CTkFrame(self.main_container)
        self.frame_logo.pack(pady=10, padx=10, fill="x")

        self.check_logo_var = ctk.StringVar(value="off")
        self.check_logo = ctk.CTkCheckBox(self.frame_logo, text="Adicionar Logomarca no Cabeçalho?", 
                                         variable=self.check_logo_var, onvalue="on", offvalue="off",
                                         command=self.toggle_logo_options, font=("Arial", 12))
        self.check_logo.pack(pady=5)

        self.scroll_logos = ctk.CTkScrollableFrame(self.frame_logo, height=80, label_text="Logos pré-definidas", label_font=("Arial", 11, "bold"))
        self.scroll_logos.pack(pady=2, padx=10, fill="x")
        
        self.logo_selecionada_var = ctk.StringVar(value="")
        self.radio_buttons_logos = []
        self.carregar_logos_da_pasta()

        # Botão e Preview (Agrupados)
        self.btn_selecionar_logo = ctk.CTkButton(self.frame_logo, text="Selecionar outra Logo", state="disabled",
                                               command=self.selecionar_logo_customizada, height=28)
        self.btn_selecionar_logo.pack(pady=2)

        self.image_preview_canvas = ctk.CTkLabel(self.frame_logo, text="Logo desativada", text_color="gray", font=("Arial", 10))
        self.image_preview_canvas.pack(pady=0) # Espaçamento zero para aproximar do alinhamento

        # Alinhamento
        self.align_var = ctk.StringVar(value="CENTER")
        self.frame_align = ctk.CTkFrame(self.frame_logo, fg_color="transparent")
        self.frame_align.pack(pady=2)
        
        opts = [("Alinhar a Esquerda", "LEFT"), ("Alinhar ao Centro", "CENTER"), ("Alinhar a Direita", "RIGHT")]
        self.radios_align = []
        for text, val in opts:
            rb = ctk.CTkRadioButton(self.frame_align, text=text, variable=self.align_var, value=val, state="disabled", width=70, font=("Arial", 11))
            rb.pack(side="left", padx=5)
            self.radios_align.append(rb)

        # --- Botão de Fotos (Cor Uniformizada) ---
        self.btn_selecionar_fotos = ctk.CTkButton(self.main_container, text="Selecionar fotos da obra", 
                                                 command=self.selecionar_fotos) # Removida cor customizada para seguir o padrão
        self.btn_selecionar_fotos.pack(pady=10)

        self.label_status = ctk.CTkLabel(self.main_container, text="Nenhuma foto selecionada...", font=("Arial", 11))
        self.label_status.pack(pady=0)

        self.btn_gerar = ctk.CTkButton(self.main_container, text="GERAR RELATÓRIO", fg_color="#27ae60", 
                                       hover_color="#2ecc71", command=self.gerar_word, height=45, font=("Arial", 13, "bold"))
        self.btn_gerar.pack(pady=15)
        # --- Rodapé da Interface Principal ---
        self.label_assinatura = ctk.CTkLabel(self.main_container, 
                                             text="Criado por: Raveli L. Araujo", 
                                             font=("Arial", 10, "italic"),
                                             text_color="gray")
        self.label_assinatura.pack(side="bottom", pady=20)

    def carregar_logos_da_pasta(self):
        for rb in self.radio_buttons_logos: rb.destroy()
        arquivos = [f for f in os.listdir(self.diretorio_logos) if f.startswith("Logo_") and f.lower().endswith(('.png', '.jpg', '.jpeg'))]
        for arquivo in arquivos:
            nome = arquivo.replace("Logo_", "").split(".")[0]
            caminho = os.path.join(self.diretorio_logos, arquivo)
            rb = ctk.CTkRadioButton(self.scroll_logos, text=nome, variable=self.logo_selecionada_var, value=caminho, 
                                    command=lambda p=caminho: self.atualizar_preview(p), font=("Arial", 11))
            rb.pack(anchor="w", pady=1)
            rb.configure(state="disabled")
            self.radio_buttons_logos.append(rb)

    def atualizar_preview(self, caminho):
        if caminho and os.path.exists(caminho):
            try:
                self.caminho_logo_final = caminho
                img = Image.open(caminho)
                img.thumbnail((120, 60)) # Preview levemente menor
                self.img_preview_referencia = ctk.CTkImage(light_image=img, dark_image=img, size=(img.width, img.height))
                self.image_preview_canvas.configure(image=self.img_preview_referencia, text="")
            except:
                self.image_preview_canvas.configure(image=None, text="Erro na imagem")
        else:
            self.image_preview_canvas.configure(image=None, text="Selecione uma logo")

    def toggle_logo_options(self):
        estado = "normal" if self.check_logo_var.get() == "on" else "disabled"
        self.btn_selecionar_logo.configure(state=estado)
        for r in self.radios_align: r.configure(state=estado)
        for rb in self.radio_buttons_logos: rb.configure(state=estado)
        
        if estado == "normal":
            self.atualizar_preview(self.logo_selecionada_var.get() if self.logo_selecionada_var.get() else None)
        else:
            self.image_preview_canvas.configure(image=None, text="Logo desativada")

    def selecionar_logo_customizada(self):
        arquivo = filedialog.askopenfilename(title="Selecione a logomarca", filetypes=[("Imagens", "*.jpg *.jpeg *.png")])
        if arquivo:
            self.logo_selecionada_var.set("custom")
            self.atualizar_preview(arquivo)

    def selecionar_fotos(self):
        arquivos = filedialog.askopenfilenames(title="Selecione as fotos da obra", filetypes=[("Imagens", "*.jpg *.jpeg *.png")])
        if arquivos:
            self.lista_fotos = list(arquivos)
            self.label_status.configure(text=f"{len(self.lista_fotos)} fotos selecionadas")

    def adicionar_num_pagina(self, parágrafo):
        parágrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = parágrafo.add_run("Página ")
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

    def gerar_word(self):
        if not self.lista_fotos:
            messagebox.showwarning("Aviso", "Selecione as fotos!")
            return
        path_save = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not path_save: return

        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.5)
            if self.check_logo_var.get() == "on" and self.caminho_logo_final:
                p = section.header.paragraphs[0]
                p.alignment = getattr(WD_ALIGN_PARAGRAPH, self.align_var.get())
                p.add_run().add_picture(self.caminho_logo_final, width=Inches(1.5))
            self.adicionar_num_pagina(section.footer.paragraphs[0])

        doc.add_heading('Relatório Fotográfico de Obra', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.add_run(f"Obra: {self.entry_obra.get()}\n").bold = True
        p.add_run(f"Responsável: {self.entry_responsavel.get()}\nData: {self.entry_data.get()}")

        table = doc.add_table(rows=0, cols=2)
        for i in range(0, len(self.lista_fotos), 2):
            row_f = table.add_row().cells
            row_l = table.add_row().cells
            for j in range(2):
                if i + j < len(self.lista_fotos):
                    para = row_f[j].paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.add_run().add_picture(self.lista_fotos[i+j], width=Inches(3.0))
                    row_l[j].text = "Legenda: ________________________________"

        try:
            doc.save(path_save)
            messagebox.showinfo("Finalizado", "Relatório gerado com sucesso!")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro: {e}")

if __name__ == "__main__":
    app = RelatorioObrasApp()
    app.mainloop()