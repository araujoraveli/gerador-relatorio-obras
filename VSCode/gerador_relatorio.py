import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os

class RelatorioObrasApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Gerador de Relatório Fotográfico - Engenharia Civil")
        self.geometry("600x550")
        ctk.set_appearance_mode("dark")

        self.lista_fotos = []

        # --- Interface ---
        self.label_titulo = ctk.CTkLabel(self, text="Relatório Fotográfico Profissional", font=("Arial", 20, "bold"))
        self.label_titulo.pack(pady=20)

        self.entry_obra = ctk.CTkEntry(self, placeholder_text="Nome da Obra / Cliente", width=400)
        self.entry_obra.pack(pady=5)

        self.entry_responsavel = ctk.CTkEntry(self, placeholder_text="Responsável Técnico (Nome/CREA)", width=400)
        self.entry_responsavel.pack(pady=5)

        self.entry_data = ctk.CTkEntry(self, placeholder_text="Data da Visita (ex: 12/03/2026)", width=400)
        self.entry_data.pack(pady=5)

        self.btn_selecionar = ctk.CTkButton(self, text="Selecionar Fotos", fg_color="#2c3e50", hover_color="#34495e", command=self.selecionar_fotos)
        self.btn_selecionar.pack(pady=20)

        self.label_status = ctk.CTkLabel(self, text="Nenhuma foto selecionada", font=("Arial", 12))
        self.label_status.pack(pady=5)

        self.btn_gerar = ctk.CTkButton(self, text="GERAR RELATÓRIO (WORD)", fg_color="#27ae60", hover_color="#2ecc71", command=self.gerar_word)
        self.btn_gerar.pack(pady=30)

    def selecionar_fotos(self):
        arquivos = filedialog.askopenfilenames(title="Selecione as fotos da obra", filetypes=[("Imagens", "*.jpg *.jpeg *.png")])
        if arquivos:
            self.lista_fotos = list(arquivos)
            self.label_status.configure(text=f"{len(self.lista_fotos)} fotos selecionadas")

    def adicionar_foto_celula(self, celula, caminho_foto):
        """Insere a foto centralizada e redimensionada na célula"""
        paragraph = celula.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        # Fixamos a largura em 3 polegadas para garantir o alinhamento lado a lado
        run.add_picture(caminho_foto, width=Inches(3.0))

    def gerar_word(self):
        if not self.lista_fotos:
            messagebox.showwarning("Aviso", "Por favor, selecione as fotos primeiro.")
            return

        path_save = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not path_save:
            return

        doc = Document()

        # Configuração de Margens Estreitas
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Título do Documento
        header = doc.add_heading('Relatório Fotográfico de Obra', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Informações de Cabeçalho
        p = doc.add_paragraph()
        p.add_run(f"Obra: {self.entry_obra.get()}\n").bold = True
        p.add_run(f"Responsável: {self.entry_responsavel.get()}\n")
        p.add_run(f"Data: {self.entry_data.get()}")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph() # Espaço

        # Criar Tabela de Fotos (2 colunas)
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False 

        num_fotos = len(self.lista_fotos)
        for i in range(0, num_fotos, 2):
            # Linha para as Imagens
            row_fotos = table.add_row().cells
            # Linha para as Legendas
            row_legendas = table.add_row().cells
            
            # Foto Esquerda
            self.adicionar_foto_celula(row_fotos[0], self.lista_fotos[i])
            row_legendas[0].text = "Legenda: ________________________________"
            
            # Foto Direita (se existir)
            if i + 1 < num_fotos:
                self.adicionar_foto_celula(row_fotos[1], self.lista_fotos[i+1])
                row_legendas[1].text = "Legenda: ________________________________"

        try:
            doc.save(path_save)
            messagebox.showinfo("Sucesso", f"Relatório gerado com sucesso em:\n{path_save}")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar: {e}")

if __name__ == "__main__":
    app = RelatorioObrasApp()
    app.mainloop()